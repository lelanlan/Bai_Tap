         
# from pymysql import connect, cursors, Error
# config = {
#         'host': 'localhost',
#         'user': 'root',
#         'password': '',
#          'database': 'blog',
#             }
# cnx = connect(**config)
# cur = cnx.cursor()
# id=39
# sql="select * from blog_uer where id="+str(id)
# cur.execute(sql)
# for row in cur:
#       print( row[0])
#       print( row[1])
#       print( row[2])

# title="123"
# contain="5778"

# print(f" UPDATE blog_uer SET title='{title}',contain='{contain}' where id = {id}")

from docx import Document
from docx.shared import Inches

document = Document()
a="lalala"

document.add_heading(a, 0)

document.add_paragraph('A plain paragraph having some ')


document.save('demo1.docx')


f=open("demo.docx")

print(f)