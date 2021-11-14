from flask import Flask, render_template, Response, request, redirect
from pymysql import connect, cursors, Error
from datetime import datetime  
from docx import Document
from docx.shared import Inches
#from app import app



config = {
    'host': 'localhost',
    'user': 'root',
    'password': '',
    'database': 'blog',
    }
cnx = connect(**config)

app = Flask(__name__)

@app.route('/index', methods=["GET"])
def render_form():
    cur = cnx.cursor()
    sql="select * from blog_uer"
    cur.execute(sql)
    #cnx.close()
    return render_template("index.html",rows=cur)


@app.route('/about')
def about():
    cur = cnx.cursor()
    sql="select * from blog_uer"
    cur.execute(sql)
    #cnx.close()
    return render_template("about.html")



@app.route('/index', methods=["POST"])
def new_post():
    cur = cnx.cursor()
    title=request.form["title"]
    contain=request.form["contain"]
    sql="INSERT INTO blog_uer (title, contain, date) VALUES (%s, %s,%s)"
    value=(title,contain,datetime.now())
    try:
        cur.execute(sql,value)
        cnx.commit()
    except:
        cnx.rollback()
    #cnx.close()
    return redirect("/index", code=302)

@app.route('/editpost/<id>', methods=["GET"])
def render_form_edit(id):
    cur = cnx.cursor()
    sql="select * from blog_uer where id="+str(id)
    cur.execute(sql)
    for row in cur:
         r0=row[0]
         r1=row[1]
         r2=row[2]


    #cnx.close()
    return render_template("editpost.html",r0=r0,r1=r1,r2=r2)



@app.route('/editpost/<id>', methods=["POST"])
def edit_post(id):
    cur = cnx.cursor()
    title=request.form["title"]
    contain=request.form["contain"]
    sql=f" UPDATE blog_uer SET title='{title}',contain='{contain}' where id = {id}"
    try:
        cur.execute(sql)
        cnx.commit()
    except:
        cnx.rollback()
    #cnx.close()
    return redirect("/editpost/"+id, code=302)


@app.route('/resignation', methods=["GET"])
def xport_file():
    #cnx.close()
    return render_template("resignation.html")

@app.route('/resignation', methods=["POST"])
def get_text():
    name=request.form["name"]
    reason=request.form["reason"]
    document = Document()
    document.add_heading(name, 0)
    document.add_paragraph(reason)
    document.save('static/demo.docx')
    return redirect("/resignation", code=302)


    

if __name__ == "__main__":
    app.run(debug=True)


